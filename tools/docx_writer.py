from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_paragraph(paragraph, text):
    """Simple parser for bold and italic markdown within a paragraph."""
    # Split text by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # It's bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italics within non-bold parts
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub_part)

def parse_table_lines(lines, start_idx):
    """Parses contiguous table lines starting from start_idx.
    Returns (headers, rows, next_idx) if valid table, else None."""
    table_lines = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            i += 1
        else:
            break
            
    if len(table_lines) < 3:  # Need at least Header, Divider, and 1 Data Row
        return None
        
    # Check if second line is a divider
    divider_line = table_lines[1]
    cells = [c.strip() for c in divider_line.split('|')[1:-1]]
    is_divider = all(re.match(r'^:?-+:?$', c) for c in cells)
    if not is_divider:
        return None
        
    headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
    rows = []
    for line in table_lines[2:]:
        row_cells = [c.strip() for c in line.split('|')[1:-1]]
        if len(row_cells) < len(headers):
            row_cells += [''] * (len(headers) - len(row_cells))
        elif len(row_cells) > len(headers):
            row_cells = row_cells[:len(headers)]
        rows.append(row_cells)
        
    return headers, rows, i

def create_docx(text_content: str, output_path: str) -> str:
    """Generates an academically formatted DOCX file with native tables support."""
    try:
        doc = Document()
        
        # Setup Academic Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
 
        lines = text_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
                
            # Check for Markdown Table
            if line.startswith('|') and line.endswith('|'):
                table_data = parse_table_lines(lines, i)
                if table_data:
                    headers, rows, next_idx = table_data
                    table = doc.add_table(rows=0, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    # Add Header Row
                    hdr_cells = table.add_row().cells
                    for col_idx, header_text in enumerate(headers):
                        p = hdr_cells[col_idx].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1.15
                        parse_markdown_to_paragraph(p, header_text)
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.bold = True
                            
                    # Add Data Rows
                    for row_data in rows:
                        row_cells = table.add_row().cells
                        for col_idx, cell_text in enumerate(row_data):
                            p = row_cells[col_idx].paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            p.paragraph_format.line_spacing = 1.15
                            parse_markdown_to_paragraph(p, cell_text)
                            for run in p.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                                
                    # Add spacing after table
                    doc.add_paragraph()
                    
                    i = next_idx
                    continue
            
            # Normal Markdown Parsing
            if line.startswith('# '):
                h = doc.add_heading(level=1)
                run = h.add_run(line[2:])
                run.font.name = 'Times New Roman'
                run.font.color.rgb = None # Reset color to default black
            elif line.startswith('## '):
                h = doc.add_heading(level=2)
                run = h.add_run(line[3:])
                run.font.name = 'Times New Roman'
                run.font.color.rgb = None
            elif line.startswith('### '):
                h = doc.add_heading(level=3)
                run = h.add_run(line[4:])
                run.font.name = 'Times New Roman'
                run.font.color.rgb = None
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                p = doc.add_paragraph(style='List Bullet')
                parse_markdown_to_paragraph(p, line[2:])
            elif re.match(r'^\d+\.\s', line):
                # Numbered list
                p = doc.add_paragraph(style='List Number')
                parse_markdown_to_paragraph(p, line[line.find(' ')+1:])
            else:
                p = doc.add_paragraph()
                parse_markdown_to_paragraph(p, line)
                
            i += 1
                
        doc.save(output_path)
        return f"Successfully saved to {output_path}"
    except Exception as e:
        return f"Error creating DOCX: {e}"
